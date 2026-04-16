# app/core/merger.py
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx.shared import Pt, Cm


class Merger:
    """
    Склейщик итогового каталога из restored_blocks.json
    """

    def __init__(
        self,
        restored_blocks_json_path: Path,
        output_docx_path: Path,
        plenary_titles: list[str] | None = None,
    ):
        self.restored_blocks_json_path = restored_blocks_json_path
        self.output_docx_path = output_docx_path
        self.plenary_titles = plenary_titles or []

    # ---------------------------------------------------------
    # public
    # ---------------------------------------------------------

    def run(self):
        docs = self._load_blocks()
        ordered_docs = self._build_order(docs)
        self._build_docx(ordered_docs)

    # ---------------------------------------------------------
    # internal
    # ---------------------------------------------------------


    def _style_run(self, run, size=12, bold=False, italic=False, superscript=False):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.superscript = superscript


    def _style_paragraph(self, p, align="justify", first_line_indent_cm=1.25, space_after_pt=6):
        if align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after_pt)
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)


    def _add_simple_paragraph(self, doc: Document, text: str, *, align="justify",
                            first_line_indent_cm=1.25, space_after_pt=6,
                            size=12, bold=False, italic=False):
        p = doc.add_paragraph()
        self._style_paragraph(
            p,
            align=align,
            first_line_indent_cm=first_line_indent_cm,
            space_after_pt=space_after_pt,
        )
        run = p.add_run(text)
        self._style_run(run, size=size, bold=bold, italic=italic)
        return p


    def _load_blocks(self) -> dict:
        with open(self.restored_blocks_json_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _build_order(self, docs: dict) -> list[dict]:
        items = list(docs.values())

        plenary = []
        regular = []

        plenary_titles_norm = {self._norm_title(x) for x in self.plenary_titles}

        for item in items:
            title = item.get("title_ru") or ""
            if self._norm_title(title) in plenary_titles_norm:
                plenary.append(item)
            else:
                regular.append(item)

        # пленарные — в порядке списка plenary_titles
        plenary_sorted = sorted(
            plenary,
            key=lambda x: self._plenary_index(x.get("title_ru") or "")
        )

        # остальные — по алфавиту
        regular_sorted = sorted(
            regular,
            key=lambda x: self._norm_title(x.get("title_ru") or "")
        )

        return plenary_sorted + regular_sorted

    def _plenary_index(self, title: str) -> int:
        norm = self._norm_title(title)
        for i, ref in enumerate(self.plenary_titles):
            if self._norm_title(ref) == norm:
                return i
        return 10**6

    def _norm_title(self, text: str) -> str:
        return " ".join(str(text).lower().split()).strip()
    
    def _append_author_line_with_superscripts(self, doc: Document, text: str, *, bold=False, italic=False):
        p = doc.add_paragraph()
        self._style_paragraph(p, align="center", first_line_indent_cm=0, space_after_pt=3)

        pattern = r"(\d{1,2}(?=(?:\(\*\))?\*?(?:,|;|\s|$)))"
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            run = p.add_run(part)
            if re.fullmatch(r"\d{1,2}", part):
                self._style_run(run, size=10, bold=bold, italic=italic, superscript=True)
            else:
                self._style_run(run, size=12 if not italic else 11, bold=bold, italic=italic)

    def _build_docx(self, docs: list[dict]):
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

        seen_titles = set()
        appended_any = False

        for item in docs:
            title = self._norm_title(item.get("title_ru") or "")
            if title and title in seen_titles:
                continue
            if title:
                seen_titles.add(title)

            if appended_any:
                doc.add_page_break()

            self._append_document(doc, item)
            appended_any = True

        self.output_docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(self.output_docx_path)

    def _append_document(self, doc: Document, item: dict):
        if item.get("udk"):
            self._add_simple_paragraph(
                doc, item["udk"],
                align="left",
                first_line_indent_cm=0,
                space_after_pt=6,
                size=12,
                bold=False,
            )

        if item.get("title_ru"):
            self._add_simple_paragraph(
                doc, item["title_ru"],
                align="center",
                first_line_indent_cm=0,
                space_after_pt=6,
                size=14,
                bold=True,
            )

        if item.get("authors_ru_block"):
            for idx, line in enumerate(item["authors_ru_block"]):
                low = line.lower()
                if "@" in line or "e-mail" in low or "email" in low or "spin" in low:
                    self._append_author_line_with_superscripts(doc, line, bold=False, italic=False)
                elif any(x in low for x in ("университет", "институт", "мгу", "мгту", "ран", "university", "institute", "faculty", "department")):
                    self._append_author_line_with_superscripts(doc, line, bold=False, italic=True)
                else:
                    self._append_author_line_with_superscripts(doc, line, bold=True, italic=False)

        if item.get("abstract_ru"):
            self._add_simple_paragraph(
                doc, "Аннотация",
                align="left",
                first_line_indent_cm=0,
                space_after_pt=0,
                size=12,
                bold=True,
            )
            self._add_simple_paragraph(
                doc, item["abstract_ru"],
                align="justify",
                first_line_indent_cm=1.25,
                space_after_pt=6,
                size=12,
                bold=False,
            )

        if item.get("keywords_ru"):
            self._add_simple_paragraph(
                doc, "Ключевые слова",
                align="left",
                first_line_indent_cm=0,
                space_after_pt=0,
                size=12,
                bold=True,
            )
            self._add_simple_paragraph(
                doc, item["keywords_ru"],
                align="justify",
                first_line_indent_cm=1.25,
                space_after_pt=6,
                size=12,
                bold=False,
            )

        has_en_block = any([
            item.get("title_en"),
            item.get("authors_en_block"),
            item.get("abstract_en"),
            item.get("keywords_en"),
        ])

        if has_en_block:
            if item.get("title_en"):
                self._add_simple_paragraph(
                    doc, item["title_en"],
                    align="center",
                    first_line_indent_cm=0,
                    space_after_pt=6,
                    size=14,
                    bold=True,
                )

            if item.get("authors_en_block"):
                for line in item["authors_en_block"]:
                    low = line.lower()
                    if "@" in line or "e-mail" in low or "email" in low or "spin" in low:
                        self._append_author_line_with_superscripts(doc, line, bold=False, italic=False)
                    elif any(x in low for x in ("university", "institute", "faculty", "department", "center", "centre")):
                        self._append_author_line_with_superscripts(doc, line, bold=False, italic=True)
                    else:
                        self._append_author_line_with_superscripts(doc, line, bold=True, italic=False)

            if item.get("abstract_en"):
                self._add_simple_paragraph(
                    doc, "Abstract",
                    align="left",
                    first_line_indent_cm=0,
                    space_after_pt=0,
                    size=12,
                    bold=True,
                )
                self._add_simple_paragraph(
                    doc, item["abstract_en"],
                    align="justify",
                    first_line_indent_cm=1.25,
                    space_after_pt=6,
                    size=12,
                    bold=False,
                )

            if item.get("keywords_en"):
                self._add_simple_paragraph(
                    doc, "Keywords",
                    align="left",
                    first_line_indent_cm=0,
                    space_after_pt=0,
                    size=12,
                    bold=True,
                )
                self._add_simple_paragraph(
                    doc, item["keywords_en"],
                    align="justify",
                    first_line_indent_cm=1.25,
                    space_after_pt=6,
                    size=12,
                    bold=False,
                )

        if item.get("reference_block"):
            self._add_simple_paragraph(
                doc, "Список литературы",
                align="left",
                first_line_indent_cm=0,
                space_after_pt=0,
                size=12,
                bold=True,
            )

            for line in item["reference_block"]:
                clean = str(line).strip()
                if not clean:
                    continue
                if "разрыв страницы" in clean.lower():
                    continue
                self._add_simple_paragraph(
                    doc, clean,
                    align="left",
                    first_line_indent_cm=0,
                    space_after_pt=3,
                    size=12,
                    bold=False,
                )

        # EN block only if there is at least title or abstract or keywords
        has_en_block = any([
            item.get("title_en"),
            item.get("abstract_en"),
            item.get("keywords_en"),
            item.get("authors_en_block"),
        ])

        if has_en_block:
            if item.get("title_en"):
                p = doc.add_paragraph()
                p.add_run(item["title_en"])

            if item.get("authors_en_block"):
                for line in item["authors_en_block"]:
                    self._append_author_line_with_superscripts(doc, line)

            if item.get("abstract_en"):
                p = doc.add_paragraph()
                p.add_run("Abstract")
                p = doc.add_paragraph()
                p.add_run(item["abstract_en"])

            if item.get("keywords_en"):
                p = doc.add_paragraph()
                p.add_run("Keywords")
                p = doc.add_paragraph()
                p.add_run(item["keywords_en"])

        if item.get("reference_block"):
            p = doc.add_paragraph()
            p.add_run("Список литературы")

            for line in item["reference_block"]:
                clean = str(line).strip()
                if not clean:
                    continue
                if "разрыв страницы" in clean.lower():
                    continue
                p = doc.add_paragraph()
                p.add_run(clean)