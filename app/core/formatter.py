# app/core/formatter.py
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


class CatalogFormatter:
    """
    Пост-форматирование собранного catalog_merged.docx
    под единый редакционный вид.
    """

    SECTION_LABELS_RU = {
        "аннотация",
        "ключевые слова",
        "введение",
        "список литературы",
        "список источников",
        "литература",
    }

    SECTION_LABELS_EN = {
        "abstract",
        "keywords",
        "references",
        "introduction",
    }

    def __init__(self, input_docx_path: Path, output_docx_path: Path):
        self.input_docx_path = input_docx_path
        self.output_docx_path = output_docx_path

        # state machine
        self.phase = "start"
        self.lang_mode = "ru"

    # ---------------------------------------------------------
    # public
    # ---------------------------------------------------------

    def run(self):
        doc = Document(self.input_docx_path)
        self._setup_document_defaults(doc)

        for p in doc.paragraphs:
            text = self._get_paragraph_text(p)

            if not text:
                self._format_empty_paragraph(p)
                continue

            # ничего не угадываем, только подчищаем мусор
            if self._is_garbage_line(text):
                p.text = ""
                continue

            # гарантируем единый шрифт и интервалы, но не меняем роли абзаца
            p.paragraph_format.space_before = Pt(0)
            if p.paragraph_format.space_after is None:
                p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                if run.font.size is None:
                    run.font.size = Pt(12)

        self.output_docx_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(self.output_docx_path)

    # ---------------------------------------------------------
    # document defaults
    # ---------------------------------------------------------

    def _setup_document_defaults(self, doc: Document):
        section = doc.sections[0]

        # Поля можно потом скорректировать под редакцию
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        normal_style.font.size = Pt(12)

    # ---------------------------------------------------------
    # classify
    # ---------------------------------------------------------

    def _is_garbage_line(self, text: str) -> bool:
        low = text.lower().strip()

        if not low:
            return True

        if "разрыв страницы" in low:
            return True

        if low in {"*", "**", "***", "-", "—"}:
            return True

        return False

    def _classify_paragraph(self, text: str) -> str:
        t = " ".join(text.split()).strip()
        low = t.lower()
        low_clean = low.replace(":", "").replace("—", "").replace("-", "").strip()

        if self._is_garbage_line(t):
            return "skip"

        if self._is_udk_line(t):
            self.phase = "after_udk"
            self.lang_mode = "ru"
            return "udk"

        # section labels
        if low.startswith("аннотация") or low_clean == "аннотация":
            self.phase = "ru_abstract"
            self.lang_mode = "ru"
            return "section_label"

        if low.startswith("ключевые слова") or low_clean == "ключевые слова":
            self.phase = "ru_keywords"
            self.lang_mode = "ru"
            return "section_label"

        if low.startswith("abstract") or low_clean == "abstract":
            self.phase = "en_abstract"
            self.lang_mode = "en"
            return "section_label"

        if low.startswith("keywords") or low_clean == "keywords":
            self.phase = "en_keywords"
            self.lang_mode = "en"
            return "section_label"

        if low_clean in {"список литературы", "список источников", "литература", "references"}:
            self.phase = "references"
            return "section_label"

        if self._is_reference_item(t):
            self.phase = "references"
            return "reference_item"

        # hard context after labels
        if self.phase == "ru_abstract":
            self.phase = "after_ru_abstract"
            return "body"

        if self.phase == "ru_keywords":
            self.phase = "after_ru_keywords"
            return "keywords_body"

        if self.phase == "en_abstract":
            self.phase = "after_en_abstract"
            return "body"

        if self.phase == "en_keywords":
            self.phase = "after_en_keywords"
            return "keywords_body"

        if self.phase == "references":
            return "reference_item"

        # contacts first
        if self._is_contact_line(t):
            return "contact"

        # context-sensitive title detection
        if self.phase in {"start", "after_udk", "after_ru_keywords", "after_en_keywords"}:
            if self._looks_like_title(t):
                # определяем язык заголовка
                if self._looks_english_title(t):
                    self.lang_mode = "en"
                    self.phase = "after_en_title"
                else:
                    self.lang_mode = "ru"
                    self.phase = "after_ru_title"
                return "title"

        # after title prefer authors
        if self.phase in {"after_ru_title", "after_en_title"} and self._looks_like_author_line(t):
            self.phase = f"after_{self.lang_mode}_authors"
            return "author"

        # after authors prefer affiliations/contact
        if self.phase in {"after_ru_authors", "after_en_authors"}:
            if self._looks_like_affiliation_line(t):
                return "affiliation"
            if self._is_contact_line(t):
                return "contact"
            if self._looks_like_author_line(t):
                return "author"

        # fallbacks
        if self._looks_like_affiliation_line(t):
            return "affiliation"

        if self._looks_like_author_line(t):
            return "author"

        if self._looks_like_keywords_content(t):
            return "keywords_body"

        return "body"
    

    def _looks_english_title(self, text: str) -> bool:
        letters = [ch for ch in text if ch.isalpha()]
        if not letters:
            return False

        latin = sum(("A" <= ch <= "Z") or ("a" <= ch <= "z") for ch in letters)
        return latin / len(letters) > 0.7

    def _is_section_label(self, text: str) -> bool:
        low = text.lower().strip()
        low = low.replace(":", "").replace("—", "").replace("-", "").strip()
        return low in self.SECTION_LABELS_RU or low in self.SECTION_LABELS_EN

    def _is_contact_line(self, text: str) -> bool:
        low = text.lower().strip()

        if self._is_email_line(text):
            return True

        if "spin" in low:
            return True

        if "e-mail" in low or "email" in low:
            return True

        if self._contains_phone(text):
            return True

        return False

    def _is_udk_line(self, text: str) -> bool:
        return text.lower().startswith("удк")

    def _is_reference_item(self, text: str) -> bool:
        t = text.strip()
        return bool(
            re.match(r"^\[\d+\]", t)
            or re.match(r"^\d+\.", t)
            or "doi" in t.lower()
            or "//" in t
        )

    def _looks_like_keywords_content(self, text: str) -> bool:
        low = text.lower().strip()

        if len(text) < 15:
            return False

        if self._is_reference_item(text):
            return False

        if low.startswith(("аннотация", "abstract", "введение", "introduction")):
            return False

        return "," in text and not low.endswith(".")

    def _looks_like_title(self, text: str) -> bool:
        t = " ".join(text.split()).strip()

        if len(t) < 12 or len(t) > 220:
            return False

        if t.endswith("."):
            return False

        if self._is_reference_item(t):
            return False

        if self._is_contact_line(t):
            return False

        if self._looks_like_affiliation_line(t):
            return False

        if self._looks_like_author_line(t):
            return False

        low = t.lower()
        bad_starts = (
            "аннотация", "abstract",
            "ключевые слова", "keywords",
            "список литературы", "список источников", "references",
            "удк", "spin"
        )
        if low.startswith(bad_starts):
            return False

        # title не должен быть похож на обычный абзац
        if t.count(",") >= 2:
            return False

        if t.count(".") >= 2:
            return False

        words = t.split()
        if len(words) < 3:
            return False

        return True

    def _looks_like_author_line(self, text: str) -> bool:
        t = " ".join(text.split()).strip()

        if not t:
            return False

        low = t.lower()
        if low.startswith(("аннотация", "abstract", "keywords", "ключевые слова", "удк")):
            return False

        if self._is_contact_line(t):
            return False

        if self._looks_like_affiliation_line(t):
            return False

        # Фамилия И.О.
        if re.search(r"\b[А-ЯЁ][а-яё-]+\s+[А-ЯЁ]\.\s*[А-ЯЁ]\.?", t):
            return True
        if re.search(r"\b[A-Z][a-z-]+\s+[A-Z]\.\s*[A-Z]\.?", t):
            return True

        # И.О. Фамилия
        if re.search(r"\b[А-ЯЁ]\.\s*[А-ЯЁ]\.\s*[А-ЯЁ][а-яё-]+", t):
            return True
        if re.search(r"\b[A-Z]\.\s*[A-Z]\.\s*[A-Z][a-z-]+", t):
            return True

        # Полное имя
        if re.search(r"\b[А-ЯЁ][а-яё-]+\s+[А-ЯЁ][а-яё-]+\s+[А-ЯЁ][а-яё-]+\b", t):
            return True
        if re.search(r"\b[A-Z][a-z-]+\s+[A-Z][a-z-]+\s+[A-Z][a-z-]+\b", t):
            return True

        # Несколько авторов через запятую
        if "," in t and (
            re.search(r"[А-ЯЁ][а-яё-]+\s+[А-ЯЁ]\.", t) or
            re.search(r"[A-Z][a-z-]+\s+[A-Z]\.", t)
        ):
            return True

        return False

    def _looks_like_affiliation_line(self, text: str) -> bool:
        t = " ".join(text.split()).strip()
        low = t.lower()

        # если это длинный обычный абзац — не affiliation
        if len(t) > 180:
            return False

        markers = (
            "университет",
            "институт",
            "мгу",
            "мгту",
            "мфти",
            "ран",
            "кафедр",
            "факульт",
            "лаборатор",
            "центр",
            "university",
            "institute",
            "department",
            "faculty",
            "laboratory",
            "centre",
            "center",
        )

        if any(m in low for m in markers):
            return True

        # строки типа "1 МГТУ..., Москва, Россия"
        if re.match(r"^\d+\s*[A-Za-zА-Яа-яЁё]", t) and "," in t:
            return True

        return False

    def _is_email_line(self, text: str) -> bool:
        return bool(re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text))

    # ---------------------------------------------------------
    # formatting
    # ---------------------------------------------------------

    def _apply_paragraph_style(self, p, kind: str, text: str):
        if kind == "skip":
            p.text = ""
            return

        # reset paragraph geometry
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if kind == "udk":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = False
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        if kind == "title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.bold = True
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
            return

        if kind == "author":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        if kind == "affiliation":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = False
                run.italic = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            return

        if kind == "contact":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = False
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
            return

        if kind == "section_label":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        if kind == "keywords_body":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.bold = False
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        if kind == "reference_item":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            for run in p.runs:
                run.bold = False
                run.italic = False
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            return

        # body
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.bold = False
            run.italic = False
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    def _format_affiliation_line(self, p, text: str):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        self._replace_paragraph_text_with_single_run(
            p,
            text,
            bold=False,
            size=11
        )

    def _format_contact_line(self, p, text: str):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        self._replace_paragraph_text_with_single_run(
            p,
            text,
            bold=False,
            size=11
        )

    def _format_empty_paragraph(self, p):
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def _format_udk(self, p):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._set_all_runs_font(p, size=12, bold=False, uppercase=False)

    def _format_title(self, p, text: str):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)

        upper_text = text
        self._replace_paragraph_text_with_single_run(
            p,
            upper_text,
            bold=True,
            size=14
        )

    def _format_author_line(self, p, text: str):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        self._replace_paragraph_text_author_aware(p, text)

    def _format_section_label(self, p):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        self._set_all_runs_font(p, size=12, bold=True, uppercase=False)

    def _format_keywords_body(self, p, text: str):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        normalized = self._normalize_keywords_text(text)

        self._replace_paragraph_text_with_single_run(
            p,
            normalized,
            bold=False,
            size=12
        )


    def _normalize_keywords_text(self, text: str) -> str:
        parts = [x.strip() for x in text.split(",")]
        cleaned = []

        for part in parts:
            if not part:
                continue

            # аббревиатуры типа GPU, CPU, SEM не трогаем
            if part.isupper():
                cleaned.append(part)
            else:
                cleaned.append(part.lower())

        return ", ".join(cleaned)

    def _format_reference_item(self, p):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_all_runs_font(p, size=12, bold=False, uppercase=False)

    def _format_body(self, p):
        self._set_base_paragraph_format(p)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._set_all_runs_font(p, size=12, bold=False, uppercase=False)

    def _set_base_paragraph_format(self, p):
        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(6)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.first_line_indent = Cm(1.25)

    def _set_all_runs_font(self, p, size=12, bold=False, uppercase=False):
        for run in p.runs:
            text = run.text.upper() if uppercase else run.text
            run.text = text
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.bold = bold

    def _clear_runs_formatting(self, p):
        # Ничего не удаляем, просто потом переприсвоим нужные свойства
        pass

    def _replace_paragraph_text_with_single_run(self, p, text: str, bold: bool, size: int):
        self._clear_paragraph_runs(p)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold

    def _replace_paragraph_text_author_aware(self, p, text: str):
        self._clear_paragraph_runs(p)

        # если есть email/телефон/коды — не делаем всю строку жирной
        if self._is_email_line(text) or self._contains_phone(text) or self._contains_code_like(text):
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
            run.bold = False
            return

        # строки с ФИО — жирным
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        if len(text) < 80:
            run.bold = True
        else:
            run.bold = False

    def _contains_phone(self, text: str) -> bool:
        return bool(re.search(r"(\+?\d[\d\-\(\)\s]{6,})", text))

    def _contains_code_like(self, text: str) -> bool:
        return bool(re.search(r"\b\d+\b", text)) and not self._looks_like_author_line(text)

    def _clear_paragraph_runs(self, p):
        for run in list(p.runs):
            run._element.getparent().remove(run._element)

    def _get_paragraph_text(self, p) -> str:
        return " ".join(p.text.replace("\xa0", " ").split()).strip()